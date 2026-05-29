"""
generate_report_docx.py  —  Robust version
Converts PROJECT_REPORT.md → AGROPORTAL_Report.docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = r"PROJECT_REPORT.md"
OUT_PATH = r"AGROPORTAL_Report.docx"

# ── hex colour constants (used in XML directly) ──────────────────────────────
HEX_GREEN_DARK  = '1B5E20'
HEX_GREEN_MID   = '2E7D32'
HEX_GREEN_LIGHT = '388E3C'
HEX_TEAL        = '00696A'
HEX_TBL_HEADER  = '1B5E20'
HEX_TBL_ALT     = 'E8F5E9'
HEX_CODE_BG     = 'F5F5F5'
HEX_GRAY        = '555555'

# ── RGBColor objects (used in font.color.rgb) ────────────────────────────────
RGB_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RGB_GREEN_DARK = RGBColor(0x1B, 0x5E, 0x20)
RGB_GREEN_MID  = RGBColor(0x2E, 0x7D, 0x32)
RGB_GREEN_LITE = RGBColor(0x38, 0x8E, 0x3C)
RGB_TEAL       = RGBColor(0x00, 0x69, 0x6A)
RGB_GRAY       = RGBColor(0x55, 0x55, 0x55)


# ── helpers ───────────────────────────────────────────────────────────────────

def xml_shading(element, fill_hex: str):
    """Attach a w:shd element with the given fill colour to a paragraph or cell tcPr."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    element.append(shd)


def set_cell_bg(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    xml_shading(tcPr, fill_hex)


def set_para_bg(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    xml_shading(pPr, fill_hex)


def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), HEX_GREEN_DARK)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def spacing(para, before=0, after=6, lines=1.15):
    pf = para.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = lines


INLINE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`([^`\n]+)`|\[([^\]]+)\]\([^\)]+\)')

def strip_md(text: str) -> str:
    text = INLINE.sub(lambda m: m.group(1) or m.group(2) or m.group(3) or m.group(4), text)
    return text.strip()


def add_inline_runs(para, text: str):
    """Parse **bold**, *italic*, `code`, [link](url) and add styled runs."""
    last = 0
    for m in INLINE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):          # **bold**
            r = para.add_run(m.group(1));  r.bold = True
        elif m.group(2):        # *italic*
            r = para.add_run(m.group(2));  r.italic = True
        elif m.group(3):        # `code`
            r = para.add_run(m.group(3))
            r.font.name      = 'Courier New'
            r.font.size      = Pt(9)
            r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        elif m.group(4):        # [link text](url)
            r = para.add_run(m.group(4));  r.underline = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def add_heading(doc, text: str, level: int):
    clean = strip_md(text)
    h = doc.add_heading(clean, level=min(level, 4))
    if not h.runs:
        return
    run = h.runs[0]
    sizes  = {1: 18, 2: 15, 3: 13, 4: 11}
    colors = {1: RGB_GREEN_DARK, 2: RGB_GREEN_MID, 3: RGB_GREEN_LITE, 4: RGB_TEAL}
    run.font.size      = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, RGB_GREEN_DARK)
    befores = {1: 18, 2: 14, 3: 10, 4: 8}
    h.paragraph_format.space_before = Pt(befores.get(level, 8))
    h.paragraph_format.space_after  = Pt(4)


def add_normal(doc, text: str, indent_cm=0.0):
    p = doc.add_paragraph()
    add_inline_runs(p, text.strip())
    spacing(p, before=2, after=5)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def add_bullet(doc, text: str, level: int = 0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())
    for r in p.runs:
        r.font.size = Pt(10)
    spacing(p, before=1, after=2, lines=1.1)


def add_numbered(doc, text: str, level: int = 0):
    style = 'List Number' if level == 0 else 'List Number 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())
    for r in p.runs:
        r.font.size = Pt(10)
    spacing(p, before=1, after=2, lines=1.1)


def add_code_block(doc, lines: list):
    for line in lines:
        p = doc.add_paragraph()
        set_para_bg(p, HEX_CODE_BG)
        r = p.add_run(line)
        r.font.name  = 'Courier New'
        r.font.size  = Pt(8.5)
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def add_blockquote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(strip_md(text.lstrip('> ').strip()))
    r.italic         = True
    r.font.size      = Pt(9.5)
    r.font.color.rgb = RGB_GRAY


def add_table(doc, header: list, rows: list):
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style         = 'Table Grid'
    table.alignment     = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True

    # header row
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEX_TBL_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(strip_md(txt))
        run.bold = True
        run.font.color.rgb = RGB_WHITE
        run.font.size      = Pt(9)

    # data rows — pad/clip to ncols
    for ridx, rdata in enumerate(rows):
        # normalise column count
        padded = list(rdata) + [''] * ncols
        padded = padded[:ncols]

        tr_cells = table.add_row().cells
        bg = HEX_TBL_ALT if ridx % 2 == 1 else None
        for i, txt in enumerate(padded):
            cell = tr_cells[i]
            if bg:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            add_inline_runs(p, strip_md(txt))
            for r in p.runs:
                r.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()   # spacing after table


def set_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)


# ── main parser ───────────────────────────────────────────────────────────────

def parse(md_path: str, out_path: str):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    i         = 0
    in_code   = False
    code_buf  = []

    # regex patterns
    re_h     = re.compile(r'^(#{1,4})\s+(.*)')
    re_hr    = re.compile(r'^(-{3,}|\*{3,})$')
    re_bq    = re.compile(r'^>')
    re_bull  = re.compile(r'^(\s*)[-*+]\s+(.*)')
    re_num   = re.compile(r'^(\s*)\d+\.\s+(.*)')
    re_tcell = re.compile(r'^\|')
    re_tsep  = re.compile(r'^[\s|:-]+$')   # separator row like |---|---|

    while i < len(lines):
        raw  = lines[i].rstrip('\n')
        line = raw

        # ── code fence ─────────────────────────────────────────────────
        if line.lstrip().startswith('```'):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── horizontal rule ────────────────────────────────────────────
        if re_hr.match(stripped):
            add_hr(doc)
            i += 1
            continue

        # ── table ──────────────────────────────────────────────────────
        if re_tcell.match(line):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                tbl_lines.append(lines[i].rstrip('\n'))
                i += 1
            parsed = []
            for tl in tbl_lines:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                # skip separator rows
                if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                    continue
                parsed.append(cells)
            if parsed:
                add_table(doc, parsed[0], parsed[1:])
            continue

        # ── heading ────────────────────────────────────────────────────
        m = re_h.match(line)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # ── blockquote ─────────────────────────────────────────────────
        if re_bq.match(line):
            add_blockquote(doc, line)
            i += 1
            continue

        # ── bullet list ────────────────────────────────────────────────
        m = re_bull.match(line)
        if m:
            lvl = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=lvl)
            i += 1
            continue

        # ── numbered list ──────────────────────────────────────────────
        m = re_num.match(line)
        if m:
            lvl = len(m.group(1)) // 2
            add_numbered(doc, m.group(2), level=lvl)
            i += 1
            continue

        # ── blank line ─────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── normal paragraph ───────────────────────────────────────────
        add_normal(doc, stripped)
        i += 1

    doc.save(out_path)
    print("Saved: " + out_path)


if __name__ == '__main__':
    parse(MD_PATH, OUT_PATH)
