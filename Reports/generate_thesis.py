import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_line(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_header_line(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'triple')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_docx(md_path, out_path):
    doc = Document()
    
    # Configure sections (Margins and Footer)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
        
        # Add footer
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Dept. of CSE, IET Lucknow\t\t\tPage No. "
        # We need a field code for page number, as a quick approach just put text, but Word needs field code.
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(10)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "AGROPORTAL: An Intelligent AI-Powered Platform"
        hp.style.font.name = 'Times New Roman'
        hp.style.font.size = Pt(10)
        add_line(hp)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print("Error reading md:", e)
        return

    in_table = False
    table_data = []

    def style_run(run, bold=False, italic=False):
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def process_text_runs(paragraph, text):
        # Basic bold processing **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = paragraph.add_run(part[2:-2])
                style_run(r, bold=True)
            else:
                r = paragraph.add_run(part)
                style_run(r)

    idx = 0
    chapter_count = 0
    while idx < len(lines):
        line = lines[idx].strip()
        idx += 1
        
        if not line:
            continue
            
        if line.startswith('---'):
            # doc.add_page_break()
            continue

        if line.startswith('# CHAPTER') or (line.startswith('#') and 'CHAPTER' in line):
            doc.add_page_break()
            chapter_count += 1
            
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r1 = p1.add_run(f"Chapter {chapter_count}")
            r1.bold = True
            r1.font.size = Pt(14)
            r1.font.name = 'Times New Roman'
            
            # The thick line
            add_header_line(p1)
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_title = line.replace('# CHAPTER', '').replace('#', '').strip()
            ch_title = ch_title.split(':')[-1].strip()
            r2 = p2.add_run(ch_title)
            r2.bold = True
            r2.font.size = Pt(16)
            r2.font.name = 'Times New Roman'
            
            p3 = doc.add_paragraph()
            add_header_line(p3)
            continue
            
        if line.startswith('#'):
            h_level = len(re.match(r'^#+', line).group(0))
            text = line.replace('#', '').strip()
            
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.name = 'Times New Roman'
            
            if h_level == 1:
                r.font.size = Pt(16)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif h_level == 2:
                r.font.size = Pt(14)
            else:
                r.font.size = Pt(13)
            continue
            
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_data = []
            cols = [c.strip() for c in line.split('|') if c.strip()]
            if not all(c.replace('-','').strip() == '' for c in cols):
                table_data.append(cols)
            continue
        elif in_table:
            # Process table
            if len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, cell_data in enumerate(row_data):
                        if c_idx < len(table.columns):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_data.replace('**', '')
                            for par in cell.paragraphs:
                                for r in par.runs:
                                    r.font.name = 'Times New Roman'
                                    r.font.size = Pt(11)
            in_table = False
            table_data = []
            
        if not in_table:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            process_text_runs(p, line)

    doc.save(out_path)
    print("Done generating document:", out_path)

if __name__ == "__main__":
    md_file = r"C:\Users\pande\Desktop\Major Project\Agriculture_Portal\PROJECT_REPORT.md"
    docx_file = r"C:\Users\pande\Desktop\Major Project\Agriculture_Portal\AGROPORTAL_THESIS.docx"
    build_docx(md_file, docx_file)
